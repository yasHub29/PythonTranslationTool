# modules/docx_translator/docx_writer.py
# 指定された位置にある段落やテーブルセル内のテキストを置換して保存する
# 書式 (最初の run のフォーマット) をできるだけ保持する戦略を採る

from docx import Document

def _replace_paragraph_text_preserve_format(para, new_text):
    """
    Paragraph のテキストを置換しつつ、最初の run のフォーマットを保持する。
    - runs がある場合、runs[0].text に new_text を入れ、他は空文字にする。
    - runs が無い場合、para.text = new_text。
    new_text に改行が含まれている場合、run.text にそのまま入れる（docx の run は改行を扱える）。
    """
    runs = para.runs
    if runs:
        # preserve formatting of runs[0]
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
    else:
        para.text = new_text

def write_docx_from_template(src_path, dest_path, translated):
    """
    src_path: 元の docx
    dest_path: 出力先
    translated: 辞書形式（reader と同じキー）で翻訳済みテキストを渡す
    """
    doc = Document(src_path)

    # paragraphs
    para_map = {i: text for (_type, i, text) in translated.get('paragraphs', []) if _type == 'para'}
    # iterate through doc.paragraphs and apply replacements where index matches
    for i, para in enumerate(doc.paragraphs):
        if i in para_map:
            new_text = para_map[i]
            try:
                _replace_paragraph_text_preserve_format(para, new_text)
            except Exception:
                # skip failures to be robust
                continue

    # tables
    # build mapping: (t_idx, r_idx, c_idx, p_idx) -> text
    table_map = {}
    for item in translated.get('tables', []):
        # ("table", t_idx, r_idx, c_idx, p_idx, text)
        _, t_idx, r_idx, c_idx, p_idx, text = item
        table_map[(t_idx, r_idx, c_idx, p_idx)] = text

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    key = (t_idx, r_idx, c_idx, p_idx)
                    if key in table_map:
                        try:
                            _replace_paragraph_text_preserve_format(para, table_map[key])
                        except Exception:
                            continue

    # headers / footers
    header_map = { (s_idx, p_idx): text for (_type, s_idx, p_idx, text) in translated.get('headers', []) if _type == 'header' }
    footer_map = { (s_idx, p_idx): text for (_type, s_idx, p_idx, text) in translated.get('footers', []) if _type == 'footer' }

    for s_idx, section in enumerate(doc.sections):
        header = section.header
        for p_idx, para in enumerate(header.paragraphs):
            key = (s_idx, p_idx)
            if key in header_map:
                try:
                    _replace_paragraph_text_preserve_format(para, header_map[key])
                except Exception:
                    continue
        footer = section.footer
        for p_idx, para in enumerate(footer.paragraphs):
            key = (s_idx, p_idx)
            if key in footer_map:
                try:
                    _replace_paragraph_text_preserve_format(para, footer_map[key])
                except Exception:
                    continue

    # save
    doc.save(dest_path)

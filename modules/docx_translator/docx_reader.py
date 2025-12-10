# modules/docx_translator/docx_reader.py
# DOCX ドキュメントから構造化テキストを抽出する
# 抽出結果は位置情報 (path) とテキストのタプルのリストとして返す

from docx import Document

def read_docx(path):
    """
    Returns a structure:
    {
      'paragraphs': [ ("para", para_index, "text"), ... ],
      'tables': [ ("table", table_index, row_idx, col_idx, para_index_within_cell, "text"), ... ],
      'headers': [ ("header", part_index, para_index, "text"), ... ],    # optional
      'footers': [ ("footer", part_index, para_index, "text"), ... ],    # optional
    }

    The path descriptors are intended for the writer to find and replace text in-place.
    Images and other non-text content are not modified and will be preserved by the writer.
    """
    doc = Document(path)
    out = {
        'paragraphs': [],
        'tables': [],
        'headers': [],
        'footers': []
    }

    # body paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text or ""
        if text.strip():
            out['paragraphs'].append(("para", i, text))

    # tables (iterate table -> row -> cell -> para)
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    text = para.text or ""
                    if text.strip():
                        out['tables'].append(("table", t_idx, r_idx, c_idx, p_idx, text))

    # headers / footers (may have multiple sections)
    for s_idx, section in enumerate(doc.sections):
        header = section.header
        for p_idx, para in enumerate(header.paragraphs):
            text = para.text or ""
            if text.strip():
                out['headers'].append(("header", s_idx, p_idx, text))
        footer = section.footer
        for p_idx, para in enumerate(footer.paragraphs):
            text = para.text or ""
            if text.strip():
                out['footers'].append(("footer", s_idx, p_idx, text))

    return out
